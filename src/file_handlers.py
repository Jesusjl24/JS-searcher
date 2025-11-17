"""
File handlers for resume processing
Handles PDF, DOCX, and TXT file formats with proper error handling
"""

import io
import logging
from typing import Optional, BinaryIO
from config import RESUME_CONFIG

# Configure logging
logger = logging.getLogger(__name__)


class UnsupportedFileTypeError(Exception):
    """Raised when an unsupported file type is encountered"""
    pass


class FileSizeExceededError(Exception):
    """Raised when file size exceeds maximum allowed"""
    pass


class FileReadError(Exception):
    """Raised when there's an error reading the file"""
    pass


def validate_file(uploaded_file, max_size_bytes: int = None) -> None:
    """
    Validate uploaded file for type and size

    Args:
        uploaded_file: Streamlit UploadedFile object
        max_size_bytes: Maximum file size in bytes

    Raises:
        UnsupportedFileTypeError: If file type is not supported
        FileSizeExceededError: If file exceeds maximum size
    """
    if max_size_bytes is None:
        max_size_bytes = RESUME_CONFIG["max_file_size_bytes"]

    # Check file extension
    file_extension = uploaded_file.name.split('.')[-1].lower()
    if file_extension not in RESUME_CONFIG["allowed_extensions"]:
        raise UnsupportedFileTypeError(
            f"File type '.{file_extension}' is not supported. "
            f"Allowed types: {', '.join(RESUME_CONFIG['allowed_extensions'])}"
        )

    # Check file size
    if uploaded_file.size > max_size_bytes:
        max_mb = max_size_bytes / (1024 * 1024)
        actual_mb = uploaded_file.size / (1024 * 1024)
        raise FileSizeExceededError(
            f"File size ({actual_mb:.2f} MB) exceeds maximum allowed size ({max_mb:.2f} MB)"
        )

    logger.info(f"File validation passed: {uploaded_file.name} ({uploaded_file.size} bytes)")


def extract_text_from_pdf(uploaded_file) -> str:
    """
    Extract text from PDF file

    Args:
        uploaded_file: Streamlit UploadedFile object or file-like object

    Returns:
        Extracted text from PDF

    Raises:
        FileReadError: If PDF cannot be read
    """
    try:
        import PyPDF2

        pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
        text = ""

        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                    logger.debug(f"Extracted {len(page_text)} chars from page {page_num + 1}")
            except Exception as e:
                logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
                continue

        # Reset file pointer for potential re-reads
        uploaded_file.seek(0)

        if not text.strip():
            raise FileReadError("No text could be extracted from PDF")

        logger.info(f"Successfully extracted {len(text)} characters from PDF")
        return text

    except ImportError:
        logger.error("PyPDF2 not installed")
        raise FileReadError("PyPDF2 library is required for PDF processing. Install with: pip install PyPDF2")
    except Exception as e:
        logger.error(f"Error reading PDF: {e}")
        raise FileReadError(f"Failed to read PDF file: {str(e)}")


def extract_text_from_docx(uploaded_file) -> str:
    """
    Extract text from DOCX file

    Args:
        uploaded_file: Streamlit UploadedFile object or file-like object

    Returns:
        Extracted text from DOCX

    Raises:
        FileReadError: If DOCX cannot be read
    """
    try:
        import docx

        doc = docx.Document(io.BytesIO(uploaded_file.read()))
        text = ""

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
            text += "\n"

        # Reset file pointer
        uploaded_file.seek(0)

        if not text.strip():
            raise FileReadError("No text could be extracted from DOCX")

        logger.info(f"Successfully extracted {len(text)} characters from DOCX")
        return text

    except ImportError:
        logger.error("python-docx not installed")
        raise FileReadError("python-docx library is required for DOCX processing. Install with: pip install python-docx")
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        raise FileReadError(f"Failed to read DOCX file: {str(e)}")


def extract_text_from_txt(uploaded_file) -> str:
    """
    Extract text from TXT file

    Args:
        uploaded_file: Streamlit UploadedFile object or file-like object

    Returns:
        Extracted text from TXT

    Raises:
        FileReadError: If TXT cannot be read
    """
    try:
        text = uploaded_file.read().decode('utf-8')

        # Reset file pointer
        uploaded_file.seek(0)

        if not text.strip():
            raise FileReadError("Text file is empty")

        logger.info(f"Successfully extracted {len(text)} characters from TXT")
        return text

    except UnicodeDecodeError:
        try:
            # Try alternative encodings
            uploaded_file.seek(0)
            text = uploaded_file.read().decode('latin-1')
            uploaded_file.seek(0)
            logger.warning("Used latin-1 encoding as fallback")
            return text
        except Exception as e:
            logger.error(f"Failed to decode text file with multiple encodings: {e}")
            raise FileReadError("Could not decode text file. File may be corrupted or use an unsupported encoding.")
    except Exception as e:
        logger.error(f"Error reading TXT: {e}")
        raise FileReadError(f"Failed to read TXT file: {str(e)}")


def extract_text_from_resume(uploaded_file) -> str:
    """
    Extract text from uploaded resume file (PDF, DOCX, or TXT)
    Main entry point for resume text extraction

    Args:
        uploaded_file: Streamlit UploadedFile object

    Returns:
        Extracted text from the resume

    Raises:
        UnsupportedFileTypeError: If file type is not supported
        FileSizeExceededError: If file exceeds maximum size
        FileReadError: If file cannot be read
    """
    # Validate file first
    validate_file(uploaded_file)

    # Determine file type and extract text
    file_extension = uploaded_file.name.split('.')[-1].lower()

    logger.info(f"Extracting text from {file_extension.upper()} file: {uploaded_file.name}")

    if file_extension == 'pdf':
        return extract_text_from_pdf(uploaded_file)
    elif file_extension == 'docx':
        return extract_text_from_docx(uploaded_file)
    elif file_extension == 'txt':
        return extract_text_from_txt(uploaded_file)
    else:
        # Should never reach here due to validation, but just in case
        raise UnsupportedFileTypeError(f"Unsupported file type: {file_extension}")


def get_file_info(uploaded_file) -> dict:
    """
    Get metadata about uploaded file

    Args:
        uploaded_file: Streamlit UploadedFile object

    Returns:
        Dictionary with file metadata
    """
    file_extension = uploaded_file.name.split('.')[-1].upper()
    file_size_kb = uploaded_file.size / 1024
    file_size_mb = uploaded_file.size / (1024 * 1024)

    return {
        "name": uploaded_file.name,
        "extension": file_extension,
        "size_bytes": uploaded_file.size,
        "size_kb": round(file_size_kb, 2),
        "size_mb": round(file_size_mb, 2),
        "is_valid": file_extension.lower() in RESUME_CONFIG["allowed_extensions"],
        "within_size_limit": uploaded_file.size <= RESUME_CONFIG["max_file_size_bytes"],
    }
